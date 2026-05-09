from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(11)

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0]
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        run.font.size = Pt(16)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        run.font.size = Pt(13)

def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()

def bullet(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

def code_block(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    doc.add_paragraph()

# ── 표지 ──
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('FoamFit 기획서')
r.bold = True; r.font.size = Pt(26)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.add_run('물티슈가 필요 없는 스마트 화장지 솔루션').italic = True

ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver_p.add_run('버전: 1.0 완성본  |  작성일: 2026-05-09').font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
doc.add_paragraph()

# ── 1. 아이디어 개요 ──
heading(doc, '1. 아이디어 개요')
p = doc.add_paragraph()
p.add_run('제품명 (가칭): ').bold = True
p.add_run('FoamFit — 폼 분사 모듈')
p = doc.add_paragraph()
p.add_run('한 줄 정의: ').bold = True
p.add_run('공중화장실 점보롤 디스펜서에 부착하는 모듈형 장치로, 화장지를 뽑는 물리적 힘만으로 기계식 펌프를 작동시켜 거품(폼) 세정제를 자동 분사함으로써 전기 없이도 물티슈 수준의 위생 경험을 제공한다.')
doc.add_paragraph()

# ── 2. 문제 정의 ──
heading(doc, '2. 문제 정의 (Pain Point)')
doc.add_paragraph('공중화장실에서 대변 후 일반 화장지만으로는 청결감이 부족하지만:')
bullet(doc, [
    '변기 막힘: 물티슈는 물에 녹지 않아 배관을 막고 수십만 원의 처리 비용 유발',
    '쓰레기 문제: 물티슈 쓰레기가 화장실 쓰레기통을 가득 채워 관리 부담 증가',
    '비데 보급 불가: 설치 비용·공간·배관 문제로 공중화장실 보급 어려움',
])
p = doc.add_paragraph()
p.add_run('→ "깨끗하게 닦고 싶지만 물티슈는 쓰기 꺼려지는" 마찰(Friction) 지점').bold = True
doc.add_paragraph()

# ── 3. 시장 검증 ──
heading(doc, '3. 시장 검증 (블라인드 앱 설문, n=18, 2026)')
table(doc,
    ['응답', '비율', '인원'],
    [
        ['물티슈가 모든 화장실에 있다면 쓴다 (= 잠재 고객)', '50.0%', '9명'],
        ['마른 휴지로도 충분하다', '33.3%', '6명'],
        ['물티슈도 귀찮아서 안 쓴다', '16.7%', '3명'],
    ]
)
p = doc.add_paragraph()
p.add_run('핵심 인사이트: ').bold = True
p.add_run('응답자 절반이 더 쾌적한 대안이 있다면 기꺼이 사용할 의향 있음.')
doc.add_paragraph()

# ── 4. 솔루션 설계 ──
heading(doc, '4. 솔루션 설계')
heading(doc, '4-1. 하드웨어 구조', 2)
code_block(doc,
    '[ 점보롤 디스펜서 ]  ───  기존 그대로 유지\n'
    '        ↓\n'
    '[ FoamFit 부착 모듈 ]\n'
    '  ├─ 기계식 펌프: 화장지 인출 시 레버 연동 → 폼 미량 분사\n'
    '  ├─ 폼 카트리지: 교체형 (OEM 폼 → 추후 자체 포뮬러)\n'
    '  ├─ 분사 노즐: 화장지 뽑히는 지점에 위치\n'
    '  └─ 사용자 Lock 스위치: 원하지 않는 사용자는 직접 끌 수 있음'
)
p = doc.add_paragraph()
p.add_run('핵심 원리: ').bold = True
p.add_run('화장지를 당기는 물리적 힘 → 기계식 펌프 압축 → 폼 분사. ')
p.add_run('전기·배터리 일절 불필요.').bold = True

heading(doc, '4-2. 주요 설계 결정', 2)
table(doc,
    ['항목', '결정 사항', '근거'],
    [
        ['분사 물질', '거품(폼) 세정제', '화장지 찢어짐 방지 + 세정력 향상'],
        ['제품 형태', '모듈형 (기존 디스펜서 부착)', 'B2B 도입 장벽 최소화'],
        ['트리거 방식', '기계식 레버 연동', '전원 불필요, 자연스러운 사용 흐름'],
        ['전원', '없음 (순수 기계식)', '설치·유지 비용 제로'],
        ['Lock 제어', '사용자 직접 온/오프', '원하지 않는 사용자 선택권 보장'],
        ['폼 공급 전략', 'OEM 시작 → 자체 개발', '빠른 출시 + 장기 기술 차별화'],
        ['파트너십 전략', '디스펜서 제조사 공동개발 우선', '호환성 해결 + 기존 유통망 활용'],
    ]
)

# ── 5. B2B 전략 ──
heading(doc, '5. B2B 전략')
heading(doc, '5-1. 타겟 고객', 2)
table(doc,
    ['타겟', '이유'],
    [
        ['고속도로 휴게소', '유동인구 최대, 빠른 사용 데이터 확보 가능'],
        ['공공기관 (정부·지자체 청사)', '공공 위생 개선 명분, 대규모 계약 가능'],
        ['대형 쇼핑몰·복합시설', '방문객 만족도 차별화 포인트'],
        ['기업 오피스 빌딩', '총무팀 의사결정 빠름, 피드백 품질 높음'],
    ]
)
heading(doc, '5-2. B2B 설득 포인트', 2)
table(doc,
    ['예상 반론', '방어 논리'],
    [
        ['"휴지 사용량 늘어 비용 증가 아닌가?"', '물티슈 변기 막힘 수리 비용 절감이 훨씬 큼'],
        ['"설비 교체 비용이 부담스럽다"', '기존 디스펜서 그대로, 모듈만 부착 (5분 설치)'],
        ['"관리가 번거롭지 않나?"', '카트리지 정기 구독 하나로 관리 일원화'],
        ['"전기 공사가 필요한가?"', '순수 기계식, 전원 일절 불필요'],
        ['"사용자가 거부감 느끼지 않나?"', 'Lock 스위치로 개인 선택권 보장'],
    ]
)

# ── 6. 비즈니스 모델 ──
heading(doc, '6. 비즈니스 모델 & 로드맵')
heading(doc, '수익 구조', 2)
table(doc,
    ['수익원', '방식', '비고'],
    [
        ['모듈 판매', '초기 하드웨어 판매', '원가 근접 가격으로 진입 유도'],
        ['폼 카트리지 구독', '월정액 or 사용량 기반', '핵심 수익'],
        ['파트너십 수익', '라이선스 or 공동 수익 배분', 'Phase 1 주요 수익'],
    ]
)
heading(doc, '3단계 로드맵', 2)
phases = [
    ('Phase 1 (0~12개월): 파트너십 공동개발', ['디스펜서 제조사 파트너십 체결', 'OEM 폼 세정제 공급 계약', 'PoC 시범 설치 (고속도로 휴게소 우선)', '기계식 펌프 프로토타입 완성']),
    ('Phase 2 (12~24개월): 카트리지 구독 정착', ['B2B 구독 런칭', '사용량 데이터 수집 및 최적화', '고객사 레퍼런스 확보']),
    ('Phase 3 (24개월+): 자체 브랜드 확장', ['자체 폼 포뮬러 개발 및 특허 출원', 'FoamFit 독립 브랜드 론칭', '직접 B2B 영업 확장']),
]
for phase_title, items in phases:
    p = doc.add_paragraph()
    p.add_run(phase_title).bold = True
    bullet(doc, items)
    doc.add_paragraph()

# ── 7. 다음 실행 과제 ──
heading(doc, '7. 다음 실행 과제')
bullet(doc, [
    '디스펜서 제조사 리스트업 및 파트너십 접촉',
    '기계식 펌프 프로토타입 제작',
    'OEM 폼 세정제 공급사 탐색 (뿌리는 비데 폼 제조사)',
    '특허 선행 조사 (기계식 펌프 + 디스펜서 연동 방식)',
    '규제·인증 검토 (식약처 세정제 관련)',
    'MVP PoC 시범 설치 장소 확보 (고속도로 휴게소 우선)',
    '블라인드 추가 설문 (표본 확대, n=100+)',
])

output = 'docs/superpowers/specs/2026-05-09-smart-toilet-paper-design.docx'
doc.save(output)
print(f'완료: {output}')
